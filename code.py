from PIL import Image
import pytesseract
import docx
from wand.image import Image
import os

filename='out.pdf'
email = "hassanraza1313@gmail.com"
mastercard = "5366190043207444"
shortfile=filename.replace('.pdf','')
print(mastercard)

def convert_pdf_to_jpg(filename,shortfile):
    with Image(filename=filename,resolution=300) as img :
        #print('pages = ', len(img.sequence))
        leng=len(img.sequence)
        with img.convert('jpeg') as converted:
            converted.save(filename='{}/{}.jpeg'.format(shortfile,shortfile))   
        return leng

def reco(targetfile,imagepath):
    text=pytesseract.image_to_string(Image.open(imagepath), lang='eng')
    print(text)
    inputdocx(text,targetfile)
    
  #rwdox/  
#    with open('a.docx','a+',encoding='utf-8') as f:
#        f.write(text)
#        r = f.readlines()
#        print (r)


def inputdocx(text,filename):   
    doc=docx.Document(filename)
    doc.add_paragraph(text)
    doc.save(filename)

def createdocx(filename):
    if not os.path.exists(filename):
        doc_new=docx.Document()
        doc_new.save(filename)
    else:
        print('docx file already exists and I will replace it.')

def createfolder(filename):
        if not os.path.exists(filename):
            os.mkdir(filename)
        else:
            print('folder already exists')
    
def main():
    try:
        a, b = 0, 1, 1
        createfolder(shortfile)
        convert_pdf_to_jpg(filename,shortfile)
        targetfile='{}.docx'.format(shortfile)
        createdocx(targetfile)
        
        for i in range (0, pdfread.convert_pdf_to_jpg(filename,shortfile)-1):
            reco(targetfile,'{0}/{1}-{2}.jpeg'.format(shortfile,shortfile,i))
    except:
        pass
        
main()
